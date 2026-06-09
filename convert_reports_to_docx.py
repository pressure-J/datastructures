from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


BASE = Path(__file__).resolve().parent


def set_run_font(run, name="宋体", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor(*color)


def set_paragraph_font(paragraph, name="宋体", size=11):
    for run in paragraph.runs:
        set_run_font(run, name=name, size=size)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 14 else WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(text)
    set_run_font(run, size=10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)

    heading_specs = {
        "Heading 1": (16, True, (31, 78, 121)),
        "Heading 2": (14, True, (44, 92, 136)),
        "Heading 3": (12.5, True, (39, 45, 52)),
        "Heading 4": (11.5, True, (39, 45, 52)),
    }
    for style_name, (size, bold, color) in heading_specs.items():
        style = doc.styles[style_name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(*color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def add_inline_markdown(paragraph, text):
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_run_font(run, name="Consolas", size=10, color=(80, 80, 80))
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            set_run_font(run, size=11, bold=True)
        else:
            run = paragraph.add_run(part)
            set_run_font(run, size=11)


def is_table_start(lines, i):
    if i + 1 >= len(lines):
        return False
    return lines[i].lstrip().startswith("|") and re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", lines[i + 1])


def parse_table(lines, i):
    table_lines = []
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        table_lines.append(lines[i].strip())
        i += 1
    rows = []
    for idx, line in enumerate(table_lines):
        if idx == 1:
            continue
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        rows.append(cells)
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c in range(cols):
            text = row[c] if c < len(row) else ""
            set_cell_text(table.cell(r, c), text, bold=(r == 0), fill=("D9EAF7" if r == 0 else None))
    doc.add_paragraph()


def add_code_block(doc, code):
    for line in code.splitlines() or [""]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=9.5, color=(40, 40, 40))


def add_image(doc, rel_path, alt):
    img_path = (BASE / rel_path).resolve()
    if not img_path.exists():
        p = doc.add_paragraph()
        add_inline_markdown(p, f"[图片缺失：{rel_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(5.9))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(alt)
    set_run_font(cap_run, name="宋体", size=9.5, color=(96, 108, 118))


def build_docx(md_path, out_path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    configure_document(doc)

    in_code = False
    code_lines = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, "\n".join(code_lines))
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if is_table_start(lines, i):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue

        image_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line.strip())
        if image_match:
            add_image(doc, image_match.group(2), image_match.group(1))
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading_match:
            level = min(len(heading_match.group(1)), 4)
            title = heading_match.group(2).strip()
            p = doc.add_paragraph(style=f"Heading {level}")
            run = p.add_run(title)
            set_run_font(run, name="黑体", size={1: 17, 2: 15, 3: 12.5, 4: 11.5}[level], bold=True)
            if level == 1:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if re.match(r"^\s*[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_inline_markdown(p, re.sub(r"^\s*[-*]\s+", "", line))
            i += 1
            continue

        if re.match(r"^\s*\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_inline_markdown(p, re.sub(r"^\s*\d+\.\s+", "", line))
            i += 1
            continue

        p = doc.add_paragraph()
        add_inline_markdown(p, line)
        i += 1

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.text = "自习室预约管理系统综合性实验报告"
        set_paragraph_font(footer, size=9)

    doc.save(out_path)


def main():
    jobs = [
        ("自习室预约管理系统分析文档.md", "自习室预约管理系统分析文档（修改版）.docx"),
        ("自习室预约管理系统设计文档.md", "自习室预约管理系统设计文档（修改版）.docx"),
    ]
    for src, dst in jobs:
        build_docx(BASE / src, BASE / dst)
        print(dst)


if __name__ == "__main__":
    main()
